from flask import Blueprint, request, jsonify
import os
import tempfile
import PyPDF2
from docx import Document as DocxDocument
import google.generativeai as genai
import chromadb
from chromadb.config import Settings
import uuid
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from src.models.user import User
from src.models.document import Document, db
from src.routes.user import token_required

# Load environment variables
load_dotenv()

document_bp = Blueprint('document', __name__)

# Configure Gemini API
genai.configure(api_key=os.getenv('GEMINI_API_KEY'))

# Initialize ChromaDB client
chroma_client = chromadb.PersistentClient(path="./chroma_db")

# Create or get collection
collection = chroma_client.get_or_create_collection(
    name="documents",
    metadata={"hnsw:space": "cosine"}
)

def get_chroma_collection():
    """Get the ChromaDB collection for documents."""
    return collection

def extract_text_from_pdf(file_path):
    """Extract text from PDF file."""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX file."""
    text = ""
    try:
        doc = DocxDocument(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
    return text

def extract_text_from_txt(file_path):
    """Extract text from TXT file."""
    text = ""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    except Exception as e:
        print(f"Error extracting text from TXT: {e}")
    return text

def chunk_text(text, chunk_size=1000, overlap=200):
    """Split text into overlapping chunks."""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start = end - overlap
        if start >= len(text):
            break
    return chunks

def generate_embeddings(text):
    """Generate embeddings using Gemini API."""
    try:
        model = 'models/text-embedding-004'
        result = genai.embed_content(
            model=model,
            content=text,
            task_type="retrieval_document"
        )
        return result['embedding']
    except Exception as e:
        print(f"Error generating embeddings: {e}")
        return None

@document_bp.route('/upload', methods=['POST'])
@token_required
def upload_document(current_user):
    """Upload and process a document."""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Save file temporarily
        filename = secure_filename(file.filename)
        temp_dir = tempfile.mkdtemp()
        file_path = os.path.join(temp_dir, filename)
        file.save(file_path)
        
        # Extract text based on file type
        file_extension = filename.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            text = extract_text_from_pdf(file_path)
        elif file_extension in ['docx', 'doc']:
            text = extract_text_from_docx(file_path)
        elif file_extension == 'txt':
            text = extract_text_from_txt(file_path)
        else:
            return jsonify({'error': 'Unsupported file type. Please upload PDF, DOCX, or TXT files.'}), 400
        
        if not text.strip():
            return jsonify({'error': 'No text could be extracted from the document'}), 400
        
        # Chunk the text
        chunks = chunk_text(text)
        
        # Generate document ID
        document_id = str(uuid.uuid4())
        
        # Create document record in database
        document = Document(
            document_id=document_id,
            filename=filename,
            file_type=file_extension,
            total_chunks=len(chunks),
            user_id=current_user.id
        )
        
        # Process each chunk
        processed_chunks = 0
        for i, chunk in enumerate(chunks):
            if chunk.strip():  # Only process non-empty chunks
                embedding = generate_embeddings(chunk)
                if embedding:
                    # Store in ChromaDB with user_id in metadata
                    chroma_collection = get_chroma_collection()
                    chroma_collection.add(
                        embeddings=[embedding],
                        documents=[chunk],
                        metadatas=[{
                            'document_id': document_id,
                            'filename': filename,
                            'chunk_index': i,
                            'chunk_count': len(chunks),
                            'user_id': current_user.id 
                        }],
                        ids=[f"{document_id}_chunk_{i}"]
                    )
                    processed_chunks += 1
        
        # Update processed chunks count
        document.chunks_processed = processed_chunks
        
        # Save document to database
        db.session.add(document)
        db.session.commit()
        
        # Clean up temporary file
        os.remove(file_path)
        os.rmdir(temp_dir)
        
        return jsonify({
            'message': 'Document uploaded and processed successfully',
            'document_id': document_id,
            'filename': filename,
            'chunks_processed': processed_chunks,
            'total_chunks': len(chunks)
        }), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Error processing document: {str(e)}'}), 500

@document_bp.route('/ask', methods=['POST'])
@token_required
def ask_question(current_user):
    """Answer a question based on uploaded documents."""
    try:
        print(f"🔍 Question received from user {current_user.id}")
        
        data = request.get_json()
        if not data or 'question' not in data:
            return jsonify({'error': 'No question provided'}), 400
        
        question = data['question']
        document_ids = data.get('document_ids', [])  # Support for multiple documents
        document_id = data.get('document_id')  # Backward compatibility
        
        print(f"📝 Question: {question}")
        print(f"📄 Document IDs: {document_ids}")
        print(f"📄 Single Document ID: {document_id}")
        
        # If single document_id is provided, add it to document_ids list
        if document_id and not document_ids:
            document_ids = [document_id]
        
        # Generate embedding for the question
        print("🧠 Generating question embedding...")
        question_embedding = generate_embeddings(question)
        if not question_embedding:
            print("❌ Failed to generate question embedding")
            return jsonify({'error': 'Failed to process question'}), 500
        
        print("✅ Question embedding generated successfully")
        
        # Search for relevant chunks
        search_kwargs = {
            'query_embeddings': [question_embedding],
            'n_results': 5
        }

        # Add user filter and optionally document filter
        if document_ids:
            # If multiple document IDs are provided, use $or operator
            if len(document_ids) > 1:
                doc_filters = [{'document_id': doc_id} for doc_id in document_ids]
                search_kwargs['where'] = {
                    '$and': [
                        {'user_id': current_user.id},
                        {'$or': doc_filters}
                    ]
                }
            # If only one document ID is provided
            elif len(document_ids) == 1:
                search_kwargs['where'] = {
                    '$and': [
                        {'user_id': current_user.id},
                        {'document_id': document_ids[0]}
                    ]
                }
            else:
                search_kwargs['where'] = {'user_id': current_user.id}
        else:
            search_kwargs['where'] = {'user_id': current_user.id}

        print(f"🔍 Search kwargs: {search_kwargs}")
        
        # Get ChromaDB collection
        collection = get_chroma_collection()
        print("📚 Querying ChromaDB...")
        results = collection.query(**search_kwargs)
        
        print(f"📊 ChromaDB results: {len(results.get('documents', []))} documents found")
        
        # If no relevant document chunks are found, answer using Gemini directly
        # We'll use a simple heuristic to check if the retrieved chunks are relevant
        # by checking if key terms from the question appear in the chunks
        
        if not results['documents'] or not results['documents'][0] or all(not chunk.strip() for chunk in results['documents'][0]):
            print("📝 No relevant chunks found, using Gemini directly...")
            model = genai.GenerativeModel('gemini-1.5-flash')
            response = model.generate_content(question)
            answer = response.text
            print("✅ Gemini response generated successfully")
            return jsonify({
                'answer': answer,
                'confidence': 'general',
                'sources': []
            }), 200
            
        # Check if the chunks are relevant to the question
        # Extract key terms from the question (words with 4+ characters)
        import re
        key_terms = [word.lower() for word in re.findall(r'\b\w{4,}\b', question.lower())]
        
        # Check if any key terms appear in the chunks
        chunks_text = ' '.join(results['documents'][0]).lower()
        relevant_terms_found = sum(1 for term in key_terms if term in chunks_text)
        
        # If less than 30% of key terms are found in the chunks, consider it irrelevant
        if len(key_terms) > 0 and relevant_terms_found / len(key_terms) < 0.3:
            model = genai.GenerativeModel('gemini-1.5-flash')
            response = model.generate_content(question)
            answer = response.text
            return jsonify({
                'answer': answer,
                'confidence': 'general',
                'sources': []
            }), 200
        
        # Prepare context from retrieved chunks
        context_chunks = results['documents'][0]
        context = "\n\n".join(context_chunks)
        
        # Generate answer using Gemini
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        prompt = f"""
        Based on the following context from uploaded documents, please answer the question.
        If the answer cannot be found in the context, please provide a general answer based on your knowledge instead of saying you don't know.
        
        Context:
        {context}
        
        Question: {question}
        
        Answer:
        """
        
        response = model.generate_content(prompt)
        answer = response.text
        
        # Prepare source information
        sources = []
        if results.get('metadatas'):
            metadatas = results['metadatas'][0] if isinstance(results['metadatas'], list) and results['metadatas'] else []
            for metadata in metadatas:
                sources.append({
                    'filename': metadata.get('filename', 'Unknown'),
                    'chunk_index': metadata.get('chunk_index', 0)
                })
        
        return jsonify({
            'answer': answer,
            'confidence': 'high' if len(context_chunks) >= 3 else 'medium',
            'sources': sources,
            'context_chunks_used': len(context_chunks)
        }), 200
        
    except Exception as e:
        return jsonify({'error': f'Error processing question: {str(e)}'}), 500

@document_bp.route('/documents', methods=['GET'])
@token_required
def list_documents(current_user):
    """List all uploaded documents for the current user."""
    try:
        # Get documents from database for current user
        documents = Document.query.filter_by(user_id=current_user.id).order_by(Document.created_at.desc()).all()
        
        return jsonify({
            'documents': [doc.to_dict() for doc in documents],
            'total_documents': len(documents)
        }), 200
        
    except Exception as e:
        return jsonify({'error': f'Error listing documents: {str(e)}'}), 500

@document_bp.route('/documents/<document_id>', methods=['DELETE'])
@token_required
def delete_document(current_user, document_id):
    """Delete a document and all its associated chunks."""
    try:
        # Find the document in database
        document = Document.query.filter_by(
            document_id=document_id,
            user_id=current_user.id
        ).first()
        
        if not document:
            return jsonify({'error': 'Document not found or access denied'}), 404
        
        # Delete all chunks from ChromaDB
        try:
            collection = get_chroma_collection()
            # Get all chunk IDs for this document
            all_results = collection.get(
                where={'document_id': document_id, 'user_id': current_user.id}
            )
            
            if all_results['ids']:
                collection.delete(ids=all_results['ids'])
                print(f"Deleted {len(all_results['ids'])} chunks from ChromaDB")
            else:
                print("No chunks found to delete from ChromaDB")
        except Exception as e:
            print(f"Warning: Could not delete from ChromaDB: {e}")
            # Continue with database deletion even if ChromaDB fails
        
        # Delete document from database
        db.session.delete(document)
        db.session.commit()
        
        print(f"Successfully deleted document {document_id} from database")
        
        return jsonify({
            'message': 'Document deleted successfully'
        }), 200
        
    except Exception as e:
        db.session.rollback()
        print(f"Error deleting document {document_id}: {e}")
        return jsonify({'error': f'Error deleting document: {str(e)}'}), 500

@document_bp.route('/documents/<document_id>', methods=['GET'])
@token_required
def get_document(current_user, document_id):
    """Get details of a specific document."""
    try:
        document = Document.query.filter_by(
            document_id=document_id,
            user_id=current_user.id
        ).first()
        
        if not document:
            return jsonify({'error': 'Document not found or access denied'}), 404
        
        return jsonify({
            'document': document.to_dict()
        }), 200
        
    except Exception as e:
        return jsonify({'error': f'Error retrieving document: {str(e)}'}), 500

@document_bp.route('/documents/<document_id>/preview', methods=['GET'])
@token_required
def preview_document(current_user, document_id):
    try:
        # Get the document
        document = Document.query.filter_by(
            document_id=document_id, 
            user_id=current_user.id
        ).first()
        
        if not document:
            return jsonify({'error': 'Document not found'}), 404
        
        # Return basic document info for preview
        preview_data = {
            'document_id': document.document_id,
            'filename': document.filename,
            'created_at': document.created_at.isoformat() if document.created_at else None,
            'file_type': document.file_type,
            'chunks_processed': document.chunks_processed,
            'total_chunks': document.total_chunks,
            'user_id': document.user_id
        }
        
        return jsonify(preview_data), 200
        
    except Exception as e:
        print(f"Error previewing document: {e}")
        return jsonify({'error': 'Failed to preview document'}), 500

@document_bp.route('/documents/<document_id>/download', methods=['GET'])
@token_required
def download_document(current_user, document_id):
    try:
        # Get the document
        document = Document.query.filter_by(
            document_id=document_id, 
            user_id=current_user.id
        ).first()
        
        if not document:
            return jsonify({'error': 'Document not found'}), 404
        
        # For now, return a simple response
        # In a real implementation, you'd stream the actual file
        return jsonify({
            'message': 'Download endpoint reached',
            'document_id': document_id,
            'filename': document.filename
        }), 200
        
    except Exception as e:
        print(f"Error downloading document: {e}")
        return jsonify({'error': 'Failed to download document'}), 500
